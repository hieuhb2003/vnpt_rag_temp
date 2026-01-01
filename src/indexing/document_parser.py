# =============================================================================
# Document Parser - Multi-format document parsing
# =============================================================================
import re
import asyncio
from abc import ABC, abstractmethod
from dataclasses import dataclass, field
from pathlib import Path
from typing import List, Dict, Any, Optional, Tuple
from uuid import uuid4

from src.models.document import Section, DocumentMetadata
from src.utils.logging import get_logger

logger = get_logger(__name__)


@dataclass
class ParsedDocument:
    """Result of document parsing."""
    title: str
    raw_content: str
    sections: List[Section] = field(default_factory=list)
    metadata: Dict[str, Any] = field(default_factory=dict)

    # Language detection result
    language: str = "en"  # "en", "vi", or "mixed"


class BaseParser(ABC):
    """Abstract base class for document parsers."""

    @abstractmethod
    async def parse(self, file_path: str, content: bytes) -> ParsedDocument:
        """
        Parse document and extract structured content.

        Args:
            file_path: Path to file (for extension detection)
            content: File content as bytes

        Returns:
            ParsedDocument with title, sections, and metadata
        """
        pass

    def _detect_language(self, text: str) -> str:
        """Detect if text is Vietnamese, English, or mixed."""
        # Vietnamese character ranges and common words
        vi_chars = set('àáảãạăằắẳẵặâầấẩẫậèéẻẽẹêềếểễệìíỉĩịòóỏõọôồốổỗộơờớởỡợùúủũụưừứửữựỳýỷỹỵđ')
        vi_words = {'và', 'hoặc', 'của', 'có', 'không', 'được', 'trong', 'với', 'là', 'cho'}

        text_lower = text.lower()
        vi_char_count = sum(1 for c in text if c in vi_chars)
        vi_word_count = sum(1 for w in text_lower.split() if w in vi_words)

        total_chars = len([c for c in text if c.isalpha()])

        if total_chars == 0:
            return "en"

        vi_ratio = (vi_char_count + vi_word_count * 5) / max(total_chars, 1)

        if vi_ratio > 0.3:
            return "vi"
        elif vi_ratio > 0.05:
            return "mixed"
        else:
            return "en"

    def _create_section_id(self) -> str:
        """Generate unique section ID."""
        return str(uuid4())

    def _create_temp_document_id(self) -> str:
        """Generate temporary document ID (will be replaced during indexing)."""
        return str(uuid4())

    def _clean_text(self, text: str) -> str:
        """Clean and normalize text."""
        # Remove excessive whitespace
        text = re.sub(r'\n\s*\n\s*\n+', '\n\n', text)
        # Normalize spaces
        text = re.sub(r' +', ' ', text)
        return text.strip()


class MarkdownParser(BaseParser):
    """Parser for Markdown files."""

    # Markdown heading patterns
    HEADING_PATTERN = re.compile(r'^(#{1,6})\s+(.+)$', re.MULTILINE)

    async def parse(self, file_path: str, content: bytes) -> ParsedDocument:
        """Parse Markdown document."""
        # Decode content
        try:
            text = content.decode('utf-8')
        except UnicodeDecodeError:
            text = content.decode('latin-1')

        # Extract title (first H1 or filename)
        title = self._extract_title(text, file_path)

        # Extract sections
        sections = self._extract_sections(text)

        # Detect language
        language = self._detect_language(text)

        return ParsedDocument(
            title=title,
            raw_content=text,
            sections=sections,
            metadata={
                "file_type": "markdown",
                "encoding": "utf-8",
            },
            language=language,
        )

    def _extract_title(self, text: str, file_path: str) -> str:
        """Extract title from first H1 or filename."""
        # Try to find first H1
        match = re.search(r'^#\s+(.+)$', text, re.MULTILINE)
        if match:
            return match.group(1).strip()

        # Fallback to filename
        return Path(file_path).stem

    def _extract_sections(self, text: str) -> List[Section]:
        """Extract sections with heading hierarchy."""
        sections = []

        # Split by headings while keeping delimiters
        parts = self.HEADING_PATTERN.split(text)

        # Reconstruct with heading info
        # parts: [preamble, '#', title1, content1, '##', title2, content2, ...]
        # If text starts with heading, preamble is empty string
        i = 0
        current_position = 0

        # Skip preamble if present (first element may not be a heading marker)
        if i < len(parts) and not parts[i].startswith('#'):
            i = 1

        while i < len(parts) - 2:
            level_markers = parts[i]
            heading = parts[i + 1].strip()
            content = parts[i + 2] if i + 2 < len(parts) else ""

            level = len(level_markers)

            # Clean content
            content = self._clean_text(content)

            sections.append(Section(
                id=self._create_section_id(),
                document_id=self._create_temp_document_id(),  # Temp ID, will be replaced
                heading=heading,
                level=level,
                section_path=str(current_position),  # Placeholder, will be regenerated
                content=content,
                summary=None,  # Will be generated by TreeBuilder
                metadata={},
                position=current_position,
            ))

            current_position += 1
            i += 3

        # Add content before first heading as intro section
        if sections and self.HEADING_PATTERN.search(text):
            first_heading_pos = text.find(sections[0].heading)
            if first_heading_pos > 0:
                intro_content = text[:first_heading_pos].strip()
                if intro_content:
                    intro = Section(
                        id=self._create_section_id(),
                        document_id=self._create_temp_document_id(),  # Temp ID, will be replaced
                        heading="Introduction",
                        level=1,
                        section_path="0",  # Placeholder, will be regenerated
                        content=intro_content,
                        summary=None,
                        metadata={},
                        position=0,
                    )
                    sections.insert(0, intro)
                    # Update positions
                    for i, sec in enumerate(sections[1:], 1):
                        sec.position = i

        return sections


class HTMLParser(BaseParser):
    """Parser for HTML files using BeautifulSoup."""

    async def parse(self, file_path: str, content: bytes) -> ParsedDocument:
        """Parse HTML document."""
        try:
            from bs4 import BeautifulSoup
        except ImportError:
            raise ImportError("BeautifulSoup4 is required for HTML parsing. Install with: pip install beautifulsoup4")

        # Parse HTML
        soup = BeautifulSoup(content, 'html.parser')

        # Extract title
        title = self._extract_title(soup, file_path)

        # Extract sections
        sections = self._extract_sections(soup)

        # Get text content for language detection
        text = soup.get_text()

        return ParsedDocument(
            title=title,
            raw_content=text,
            sections=sections,
            metadata={
                "file_type": "html",
                "encoding": "utf-8",
            },
            language=self._detect_language(text),
        )

    def _extract_title(self, soup, file_path: str) -> str:
        """Extract title from title tag or first H1."""
        # Try title tag
        title_tag = soup.find('title')
        if title_tag and title_tag.string:
            return title_tag.string.strip()

        # Try first H1
        h1 = soup.find('h1')
        if h1:
            return h1.get_text().strip()

        # Fallback to filename
        return Path(file_path).stem

    def _extract_sections(self, soup) -> List[Section]:
        """Extract sections from heading tags."""
        sections = []

        # Find all headings (h1-h6)
        headings = soup.find_all(['h1', 'h2', 'h3', 'h4', 'h5', 'h6'])

        if not headings:
            # No headings found, create single section
            content = soup.get_text()
            content = self._clean_text(content)
            if content:
                sections.append(Section(
                    id=self._create_section_id(),
                    document_id=self._create_temp_document_id(),  # Temp ID
                    heading="Content",
                    level=1,
                    section_path="0",  # Placeholder
                    content=content,
                    summary=None,
                    metadata={},
                    position=0,
                ))
            return sections

        for i, heading in enumerate(headings):
            level = int(heading.name[1])  # h1 -> 1, h2 -> 2, etc.

            # Get heading text
            heading_text = heading.get_text().strip()

            # Get content until next heading
            content_parts = []
            next_element = heading.next_sibling

            while next_element:
                if next_element.name in ['h1', 'h2', 'h3', 'h4', 'h5', 'h6']:
                    break
                if hasattr(next_element, 'get_text'):
                    text = next_element.get_text()
                    if text.strip():
                        content_parts.append(text.strip())
                next_element = next_element.next_sibling

            content = '\n'.join(content_parts)
            content = self._clean_text(content)

            sections.append(Section(
                id=self._create_section_id(),
                document_id=self._create_temp_document_id(),  # Temp ID
                heading=heading_text,
                level=level,
                section_path=str(i),  # Placeholder
                content=content,
                summary=None,
                metadata={},
                position=i,
            ))

        return sections


class PDFParser(BaseParser):
    """Parser for PDF files using PyMuPDF (fitz)."""

    async def parse(self, file_path: str, content: bytes) -> ParsedDocument:
        """Parse PDF document."""
        try:
            import fitz  # PyMuPDF
        except ImportError:
            raise ImportError("PyMuPDF is required for PDF parsing. Install with: pip install pymupdf")

        # Open PDF from bytes
        doc = fitz.open(stream=content, filetype="pdf")

        # Extract text and metadata
        text_parts = []
        for page in doc:
            text_parts.append(page.get_text())

        full_text = '\n'.join(text_parts)

        # Extract title
        title = self._extract_title(doc, file_path)

        # Extract sections
        sections = self._extract_sections(full_text)

        # Get PDF metadata
        pdf_metadata = doc.metadata if hasattr(doc, 'metadata') else {}

        doc.close()

        return ParsedDocument(
            title=title,
            raw_content=full_text,
            sections=sections,
            metadata={
                "file_type": "pdf",
                "pages": len(doc),
                "author": pdf_metadata.get('author', ''),
                "subject": pdf_metadata.get('subject', ''),
                "keywords": pdf_metadata.get('keywords', ''),
            },
            language=self._detect_language(full_text),
        )

    def _extract_title(self, doc, file_path: str) -> str:
        """Extract title from PDF metadata or filename."""
        metadata = doc.metadata if hasattr(doc, 'metadata') else {}

        # Try metadata title
        if metadata.get('title'):
            return metadata['title'].strip()

        # Fallback to filename
        return Path(file_path).stem

    def _extract_sections(self, text: str) -> List[Section]:
        """
        Extract sections from PDF text.
        PDFs don't have explicit structure, so we try to detect headings
        by looking for all-caps or bold-like patterns (larger text usually).
        """
        sections = []

        # Split into paragraphs
        paragraphs = text.split('\n\n')

        current_section = None
        current_content = []
        position = 0

        for para in paragraphs:
            para = para.strip()
            if not para:
                continue

            # Detect potential heading (all caps or short + next para is long)
            is_heading = (
                para.isupper() and len(para) < 100 or
                (len(para) < 80 and para[0].isupper())
            )

            if is_heading and not para.isupper():
                # Skip if it's just a short sentence
                if len(para.split()) < 3:
                    is_heading = False

            if is_heading:
                # Save previous section
                if current_section:
                    current_section.content = self._clean_text('\n'.join(current_content))
                    sections.append(current_section)

                # Start new section
                current_section = Section(
                    id=self._create_section_id(),
                    document_id=self._create_temp_document_id(),  # Temp ID
                    heading=para,
                    level=1,  # Default to level 1 for PDF
                    section_path=str(position),  # Placeholder
                    content='',
                    summary=None,
                    metadata={},
                    position=position,
                )
                current_content = []
                position += 1
            else:
                if not current_section:
                    # Create intro section
                    current_section = Section(
                        id=self._create_section_id(),
                        document_id=self._create_temp_document_id(),  # Temp ID
                        heading="Introduction",
                        level=1,
                        section_path="0",  # Placeholder
                        content='',
                        summary=None,
                        metadata={},
                        position=position,
                    )
                    position += 1

                current_content.append(para)

        # Add last section
        if current_section:
            current_section.content = self._clean_text('\n'.join(current_content))
            sections.append(current_section)

        # If no sections found, create one
        if not sections and text.strip():
            sections.append(Section(
                id=self._create_section_id(),
                document_id=self._create_temp_document_id(),  # Temp ID
                heading="Content",
                level=1,
                section_path="0",  # Placeholder
                content=self._clean_text(text),
                summary=None,
                metadata={},
                position=0,
            ))

        return sections


class DocxParser(BaseParser):
    """Parser for DOCX files using python-docx."""

    async def parse(self, file_path: str, content: bytes) -> ParsedDocument:
        """Parse DOCX document."""
        try:
            import docx
        except ImportError:
            raise ImportError("python-docx is required for DOCX parsing. Install with: pip install python-docx")

        # Run in executor as python-docx is synchronous
        def parse_sync():
            # Create temp file for docx to read
            import tempfile
            import os

            with tempfile.NamedTemporaryFile(delete=False, suffix='.docx') as tmp:
                tmp.write(content)
                tmp_path = tmp.name

            try:
                doc = docx.Document(tmp_path)

                # Extract title
                title = self._extract_title(doc, file_path)

                # Extract sections
                sections = self._extract_sections(doc)

                # Get full text
                full_text = '\n'.join([p.text for p in doc.paragraphs])

                # Get core properties
                props = doc.core_properties

                return ParsedDocument(
                    title=title,
                    raw_content=full_text,
                    sections=sections,
                    metadata={
                        "file_type": "docx",
                        "author": props.author or '',
                        "subject": props.subject or '',
                        "keywords": props.keywords or '',
                        "comments": props.comments or '',
                    },
                    language=self._detect_language(full_text),
                )
            finally:
                os.unlink(tmp_path)

        return await asyncio.get_event_loop().run_in_executor(None, parse_sync)

    def _extract_title(self, doc, file_path: str) -> str:
        """Extract title from document properties or first heading."""
        # Try core properties
        if doc.core_properties.title:
            return doc.core_properties.title.strip()

        # Try first paragraph (often title)
        if doc.paragraphs and doc.paragraphs[0].text.strip():
            first_para = doc.paragraphs[0].text.strip()
            if len(first_para) < 100:  # Reasonable title length
                return first_para

        # Fallback to filename
        return Path(file_path).stem

    def _extract_sections(self, doc) -> List[Section]:
        """Extract sections from heading styles."""
        sections = []
        current_section = None
        current_content = []
        position = 0

        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                continue

            # Check if paragraph is a heading (based on style)
            style_name = para.style.name if para.style else ""

            if style_name.startswith('Heading'):
                # Extract heading level
                try:
                    level = int(style_name[-1])  # "Heading 1" -> 1
                except (ValueError, IndexError):
                    level = 1

                # Save previous section
                if current_section and current_content:
                    current_section.content = self._clean_text('\n'.join(current_content))
                    sections.append(current_section)

                # Create new section
                current_section = Section(
                    id=self._create_section_id(),
                    document_id=self._create_temp_document_id(),  # Temp ID
                    heading=text,
                    level=level,
                    section_path=str(position),  # Placeholder
                    content='',
                    summary=None,
                    metadata={},
                    position=position,
                )
                current_content = []
                position += 1

            elif current_section:
                # Add to current section
                current_content.append(text)
            else:
                # Create intro section
                current_section = Section(
                    id=self._create_section_id(),
                    document_id=self._create_temp_document_id(),  # Temp ID
                    heading="Introduction",
                    level=1,
                    section_path="0",  # Placeholder
                    content='',
                    summary=None,
                    metadata={},
                    position=0,
                )
                current_content.append(text)
                position = 1

        # Add last section
        if current_section and current_content:
            current_section.content = self._clean_text('\n'.join(current_content))
            sections.append(current_section)

        # If no sections found, create one
        if not sections:
            full_text = '\n'.join([p.text for p in doc.paragraphs])
            sections.append(Section(
                id=self._create_section_id(),
                document_id=self._create_temp_document_id(),  # Temp ID
                heading="Content",
                level=1,
                section_path="0",  # Placeholder
                content=self._clean_text(full_text),
                summary=None,
                metadata={},
                position=0,
            ))

        return sections


class DocumentParserFactory:
    """Factory for getting appropriate parser by file extension."""

    _parsers = {
        '.md': MarkdownParser(),
        '.markdown': MarkdownParser(),
        '.html': HTMLParser(),
        '.htm': HTMLParser(),
        '.pdf': PDFParser(),
        '.docx': DocxParser(),
        '.doc': DocxParser(),  # Will try to parse .doc as .docx
    }

    @classmethod
    async def parse(cls, file_path: str, content: bytes) -> ParsedDocument:
        """
        Parse document using appropriate parser.

        Args:
            file_path: Path to file
            content: File content as bytes

        Returns:
            ParsedDocument

        Raises:
            ValueError: If file type is not supported
        """
        ext = Path(file_path).suffix.lower()

        parser = cls._parsers.get(ext)
        if not parser:
            raise ValueError(
                f"Unsupported file type: {ext}. "
                f"Supported types: {', '.join(cls._parsers.keys())}"
            )

        logger.info(f"Parsing {file_path} with {parser.__class__.__name__}")

        try:
            return await parser.parse(file_path, content)
        except Exception as e:
            logger.error(f"Failed to parse {file_path}: {e}")
            raise

    @classmethod
    def is_supported(cls, file_path: str) -> bool:
        """Check if file type is supported."""
        ext = Path(file_path).suffix.lower()
        return ext in cls._parsers

    @classmethod
    def supported_extensions(cls) -> List[str]:
        """Get list of supported file extensions."""
        return list(cls._parsers.keys())


# Singleton instance for easy access
parser_factory = DocumentParserFactory
